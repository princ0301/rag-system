from docx import Document
from app.parsers.base_parser import BaseParser, ParsedDocument

class DOCXParser(BaseParser):

    def parse(self, source: str) -> ParsedDocument:
        doc = Document(source)
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return ParsedDocument(
            content=content.strip(),
            metadata={"source": source, "type": "docx"}
        )